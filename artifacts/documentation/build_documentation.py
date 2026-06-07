from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("KE03_INTDEV_SE_2_documentatie.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 107, 122)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table_header(row):
    for cell in row.cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Documentatie KE03_INTDEV_SE_2")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run("ASP.NET Core MVC admin back-office voor Matrix Inc.")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = MUTED


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)
    return paragraph


def add_h1(doc, text):
    paragraph = doc.add_heading(text, level=1)
    paragraph.runs[0].font.color.rgb = BLUE
    return paragraph


def add_h2(doc, text):
    paragraph = doc.add_heading(text, level=2)
    paragraph.runs[0].font.color.rgb = BLUE
    return paragraph


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [16.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = DARK_BLUE
    paragraph.add_run(text)
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [4.8, 11.7])
    header = table.rows[0].cells
    header[0].text = "Onderdeel"
    header[1].text = "Beschrijving"
    style_table_header(table.rows[0])
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_width(table, [4.8, 11.7])
    doc.add_paragraph()
    return table


def add_file_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [5.4, 4.5, 6.6])
    header = table.rows[0].cells
    header[0].text = "Bestand"
    header[1].text = "Type"
    header[2].text = "Doel"
    style_table_header(table.rows[0])
    for file_name, file_type, goal in rows:
        cells = table.add_row().cells
        cells[0].text = file_name
        cells[1].text = file_type
        cells[2].text = goal
    set_table_width(table, [5.4, 4.5, 6.6])
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    add_title(doc)

    add_h1(doc, "1. Doel van het project")
    doc.add_paragraph(
        "Deze documentatie beschrijft de admin back-office functionaliteit voor Matrix Inc. "
        "De applicatie is gebouwd met ASP.NET Core MVC en gebruikt de bestaande Data Access Layer. "
        "De nadruk ligt op productbeheer, orderhistorie, duidelijke navigatie en eenvoudige HCI-principes."
    )

    add_note(
        doc,
        "Kern",
        "De oplossing gebruikt Controllers, Views en ViewModels. Er zijn geen Razor Pages toegevoegd.",
    )

    add_h1(doc, "2. Technische basis")
    add_key_value_table(
        doc,
        [
            ("Framework", "ASP.NET Core MVC met AddControllersWithViews en MapControllerRoute in Program.cs."),
            ("Data Access Layer", "MatrixIncDbContext met DbSet voor Customers, Orders, Products en Parts."),
            ("Database", "SQLite via Data Source=MatrixInc.db."),
            ("Repositories", "Dependency injection voor onder andere IProductRepository, IOrderRepository en ICustomerRepository."),
            ("Routing", "Conventionele MVC-route: {controller=Home}/{action=Index}/{id?}."),
        ],
    )

    add_h1(doc, "3. Productbeheer")
    doc.add_paragraph(
        "Voor productbeheer is een ProductController toegevoegd. Deze controller gebruikt IProductRepository "
        "via dependency injection. Omdat de bestaande repository synchronische methoden bevat, zijn de acties "
        "synchronisch opgezet."
    )

    add_h2(doc, "CRUD-functionaliteit")
    for item in [
        "Index toont alle producten in een overzichtelijke tabel.",
        "Details toont de gegevens van een specifiek product.",
        "Create bevat een formulier om een product toe te voegen.",
        "Edit bevat een formulier om een bestaand product aan te passen.",
        "Delete toont eerst een bevestigingspagina voordat het product definitief wordt verwijderd.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "Zoeken en filteren")
    doc.add_paragraph(
        "De Index-actie ondersteunt zoeken op productnaam met de parameter searchString. "
        "Daarnaast kan lowStockOnly worden gebruikt om alleen producten met lage voorraad te tonen."
    )
    add_key_value_table(
        doc,
        [
            ("Zoeken", "Producten worden gefilterd op Name. De zoekactie is hoofdletterongevoelig."),
            ("Lage voorraad", "Een product heeft lage voorraad wanneer Stock kleiner is dan 5."),
            ("Reset", "De overzichtspagina heeft een resetknop om alle filters te verwijderen."),
        ],
    )

    add_h2(doc, "Validatie")
    for item in [
        "Productnaam is verplicht en maximaal 100 tekens.",
        "Omschrijving is verplicht en maximaal 500 tekens.",
        "Prijs moet groter zijn dan 0.",
        "Voorraad mag niet negatief zijn.",
        "Foutmeldingen worden direct bij de invoervelden getoond.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "4. Admin dashboard")
    doc.add_paragraph(
        "Het admin dashboard geeft snel inzicht in de productstatus. De pagina is bereikbaar via /Admin "
        "en bevat acties om naar het productoverzicht of het formulier voor een nieuw product te gaan."
    )
    add_key_value_table(
        doc,
        [
            ("Totaal aantal producten", "Aantal records uit IProductRepository.GetAllProducts()."),
            ("Lage voorraad", "Aantal producten waarbij Stock kleiner is dan 5."),
            ("Actieknoppen", "Knoppen naar Producten bekijken en Nieuw product."),
        ],
    )

    add_h1(doc, "5. Orderhistorie")
    doc.add_paragraph(
        "Naast productbeheer is een orderhistoriepagina toegevoegd. Deze pagina toont orders van alle klanten "
        "en ondersteunt filters op klant, zoekterm en datumbereik."
    )
    for item in [
        "Route: /Orders.",
        "Data wordt geladen via MatrixIncDbContext met Customer en Products.",
        "De tabel toont ordernummer, datum, klant, adres, klantstatus en producten.",
        "Bij geen resultaten wordt een duidelijke melding getoond.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "6. Homepagina en navigatie")
    doc.add_paragraph(
        "De standaard homepagina is vervangen door een duidelijke back-office startpagina. "
        "Deze pagina helpt gebruikers direct naar de belangrijkste onderdelen van de applicatie te navigeren."
    )
    add_key_value_table(
        doc,
        [
            ("Hero-sectie", "Introductie van de Matrix Inc. admin back-office met knoppen naar dashboard en productbeheer."),
            ("Navigatiekaarten", "Kaarten naar Admin dashboard, Productbeheer, Orderhistorie en Klantenbeheer."),
            ("Workflow-blok", "Korte uitleg van een logische beheerworkflow: dashboard, productbeheer en orderhistorie."),
            ("Responsive CSS", "De kaarten, knoppen en hero-sectie schalen mee naar kleinere schermen."),
        ],
    )

    add_h1(doc, "7. Figma interface design")
    doc.add_paragraph(
        "Voor de interface is een Figma-ready designpakket gemaakt in C:\\documentatie\\figma-design. "
        "Het pakket bevat importeerbare SVG-schermen, design tokens en een Nederlandse onderbouwing."
    )
    add_key_value_table(
        doc,
        [
            ("matrix-inc-backoffice-figma.svg", "Hoofdcanvas met alle schermen en componenten om in Figma te importeren."),
            ("Losse SVG-schermen", "Home, Admin dashboard, Productbeheer, Orderhistorie en Componenten als aparte imports."),
            ("design-tokens.json", "Kleuren, typografie, spacing, radius en HCI-principes."),
            ("Matrix_Inc_Figma_Interface_Design.docx", "Onderbouwing van het interface design in duidelijk Nederlands."),
        ],
    )

    add_h1(doc, "8. Bestandenoverzicht")
    add_file_table(
        doc,
        [
            ("Views/Home/Index.cshtml", "View", "Nieuwe startpagina met navigatiekaarten en snelle acties."),
            ("Controllers/ProductController.cs", "Controller", "CRUD-acties, zoeken en lage-voorraadfilter voor producten."),
            ("Controllers/AdminController.cs", "Controller", "Dashboardstatistieken voor productbeheer."),
            ("Views/Product/*.cshtml", "Views", "Overzicht, details, create, edit, delete en gedeeld formulier."),
            ("Views/Admin/Index.cshtml", "View", "Admin dashboard met statistieken en actieknoppen."),
            ("ViewModels/ProductIndexViewModel.cs", "ViewModel", "Data voor productoverzicht en filters."),
            ("ViewModels/ProductFormViewModel.cs", "ViewModel", "Formulierdata en validatie voor create/edit."),
            ("ViewModels/AdminDashboardViewModel.cs", "ViewModel", "Dashboardtellingen."),
            ("Models/Product.cs", "Model", "Productdata met Name, Description, Price en Stock."),
            ("MatrixIncDbInitializer.cs", "Data seed/schema", "Seed-voorraad en schema-check voor Stock in SQLite."),
            ("Views/Shared/_Layout.cshtml", "Layout", "Navigatie naar Admin, Products en Orders."),
            ("wwwroot/css/site.css", "CSS", "Responsive styling voor dashboard, formulieren en tabellen."),
        ],
    )

    add_h1(doc, "9. HCI-principes")
    for item in [
        "De homepagina biedt duidelijke startpunten voor de belangrijkste taken.",
        "Duidelijke primaire knoppen voor toevoegen, opslaan en filteren.",
        "Secundaire knoppen voor annuleren, resetten en teruggaan.",
        "Overzichtelijke tabellen met duidelijke kolommen en acties per rij.",
        "Bevestigingspagina voordat een product wordt verwijderd.",
        "Validatiemeldingen bij verkeerde invoer.",
        "Responsive layout zodat filters, knoppen en tabellen ook op kleinere schermen bruikbaar blijven.",
        "Logische navigatie via de hoofdnavbar.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "10. Routes om te testen")
    add_key_value_table(
        doc,
        [
            ("/", "Homepagina met navigatie naar de back-office onderdelen."),
            ("/Admin", "Admin dashboard."),
            ("/Product", "Productoverzicht met zoeken en lage-voorraadfilter."),
            ("/Product/Create", "Nieuw product toevoegen."),
            ("/Product/Details/{id}", "Details van een product bekijken."),
            ("/Product/Edit/{id}", "Product wijzigen."),
            ("/Product/Delete/{id}", "Product verwijderen na bevestiging."),
            ("/Orders", "Orderhistorie van alle klanten."),
        ],
    )

    add_h1(doc, "11. Controle en build")
    doc.add_paragraph(
        "De oplossing is gecontroleerd met dotnet build. De build is geslaagd. Er zijn nog bestaande nullable "
        "warnings in Customer en Part, maar geen compile errors in de toegevoegde admin- en productbeheerfunctionaliteit."
    )
    add_numbered(doc, "Start de applicatie vanuit Visual Studio of met dotnet run.")
    add_numbered(doc, "Open /Admin voor het dashboard.")
    add_numbered(doc, "Open /Product om productbeheer te testen.")
    add_numbered(doc, "Test create, edit, details en delete met geldige en ongeldige invoer.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("KE03_INTDEV_SE_2 - Matrix Inc. documentatie")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
