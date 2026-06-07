import json
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent
CANVAS = OUT / "matrix-inc-backoffice-figma.svg"
TOKENS = OUT / "design-tokens.json"
GUIDE = OUT / "figma-import-guide.md"
RATIONALE_MD = OUT / "interface-design-onderbouwing.md"
RATIONALE_DOCX = OUT / "Matrix_Inc_Figma_Interface_Design.docx"


COLORS = {
    "ink": "#0B2545",
    "heading": "#1F4D78",
    "primary": "#2E74B5",
    "primary_dark": "#1F5E93",
    "success": "#198754",
    "warning": "#FFD76A",
    "danger": "#DC3545",
    "muted": "#5F6B7A",
    "border": "#D9E0E8",
    "surface": "#FFFFFF",
    "surface_alt": "#F7F9FB",
    "hero": "#EEF5FB",
    "page": "#F4F7FA",
}


TOKENS_DATA = {
    "name": "Matrix Inc Back-office UI",
    "colors": COLORS,
    "typography": {
        "fontFamily": "Inter, Segoe UI, Arial, sans-serif",
        "h1": {"size": 36, "weight": 700, "lineHeight": 44},
        "h2": {"size": 24, "weight": 700, "lineHeight": 32},
        "h3": {"size": 18, "weight": 700, "lineHeight": 26},
        "body": {"size": 16, "weight": 400, "lineHeight": 24},
        "small": {"size": 13, "weight": 500, "lineHeight": 18},
    },
    "spacing": {"xs": 4, "sm": 8, "md": 16, "lg": 24, "xl": 32, "xxl": 48},
    "radius": {"sm": 6, "md": 8},
    "breakpoints": {"mobile": 375, "tablet": 768, "desktop": 1366},
    "hciPrinciples": [
        "Duidelijke navigatie vanaf de homepagina.",
        "Primaire acties zijn visueel sterker dan secundaire acties.",
        "Tabellen zijn scanbaar door kolomkoppen, badges en consistente actieknoppen.",
        "Verwijderen gebeurt via een aparte bevestigingsstap.",
        "Filters staan boven het overzicht en behouden de taakcontext.",
    ],
}


def rect(x, y, w, h, fill, stroke=None, radius=8, sw=1, extra=""):
    stroke_attr = f' stroke="{stroke}" stroke-width="{sw}"' if stroke else ""
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{radius}" fill="{fill}"{stroke_attr} {extra}/>'


def text(x, y, content, size=16, weight=400, color=None, anchor="start"):
    color = color or COLORS["ink"]
    return (
        f'<text x="{x}" y="{y}" font-family="Inter, Segoe UI, Arial, sans-serif" '
        f'font-size="{size}" font-weight="{weight}" fill="{color}" text-anchor="{anchor}">'
        f"{escape(content)}</text>"
    )


def line(x1, y1, x2, y2, color=None, sw=1):
    color = color or COLORS["border"]
    return f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="{sw}"/>'


def button(x, y, label, variant="primary", width=150):
    if variant == "primary":
        fill, stroke, color = COLORS["primary"], None, "#FFFFFF"
    elif variant == "danger":
        fill, stroke, color = COLORS["danger"], None, "#FFFFFF"
    elif variant == "success":
        fill, stroke, color = COLORS["success"], None, "#FFFFFF"
    else:
        fill, stroke, color = "#FFFFFF", COLORS["border"], COLORS["heading"]
    return [
        rect(x, y, width, 42, fill, stroke, 6),
        text(x + width / 2, y + 27, label, 14, 700, color, "middle"),
    ]


def navbar(x, y, active):
    items = ["Home", "Customers", "Admin", "Products", "Privacy", "Orders"]
    parts = [
        rect(x, y, 1366, 64, COLORS["surface"], COLORS["border"], 0),
        text(x + 32, y + 39, "Matrix Inc.", 20, 700, COLORS["ink"]),
    ]
    cursor = x + 205
    for item in items:
        is_active = item == active
        if is_active:
            parts.append(rect(cursor - 10, y + 15, len(item) * 10 + 24, 34, COLORS["hero"], None, 6))
        parts.append(text(cursor, y + 38, item, 14, 700 if is_active else 500, COLORS["heading"] if is_active else COLORS["muted"]))
        cursor += len(item) * 10 + 44
    return parts


def frame_label(x, y, title, subtitle):
    return [
        text(x, y - 18, title, 18, 800, COLORS["ink"]),
        text(x, y + 4, subtitle, 13, 500, COLORS["muted"]),
    ]


def shell(x, y, title, subtitle, active):
    return [
        rect(x, y, 1366, 900, COLORS["page"], COLORS["border"], 8),
        *navbar(x, y, active),
        text(x + 48, y + 126, title, 34, 800, COLORS["ink"]),
        text(x + 48, y + 158, subtitle, 16, 400, COLORS["muted"]),
    ]


def home_frame(x, y):
    parts = [*frame_label(x, y - 22, "01 Home", "Startpunt voor alle beheertaken"), *shell(x, y, "Back-office overzicht", "Navigeer snel naar de belangrijkste onderdelen.", "Home")]
    parts += [
        rect(x + 48, y + 198, 840, 250, COLORS["hero"], COLORS["border"], 8),
        text(x + 80, y + 250, "Matrix Inc. Admin", 14, 800, COLORS["heading"]),
        text(x + 80, y + 302, "Rustige startpagina voor dagelijks beheer", 30, 800, COLORS["ink"]),
        text(x + 80, y + 346, "Primaire routes staan vooraan. De gebruiker hoeft niet te zoeken in menu's.", 16, 400, COLORS["muted"]),
        *button(x + 80, y + 382, "Open dashboard", "primary", 172),
        *button(x + 266, y + 382, "Productbeheer", "outline", 160),
        rect(x + 936, y + 198, 300, 250, COLORS["surface"], COLORS["border"], 8),
        text(x + 966, y + 246, "Producten", 14, 500, COLORS["muted"]),
        text(x + 966, y + 276, "CRUD", 24, 800, COLORS["ink"]),
        line(x + 966, y + 302, x + 1204, y + 302),
        text(x + 966, y + 338, "Orders", 14, 500, COLORS["muted"]),
        text(x + 966, y + 368, "Historie", 24, 800, COLORS["ink"]),
        line(x + 966, y + 394, x + 1204, y + 394),
        text(x + 966, y + 430, "Voorraad", 14, 500, COLORS["muted"]),
        text(x + 966, y + 460, "Signalering", 24, 800, COLORS["ink"]),
    ]
    cards = [
        ("Dashboard", "Admin overzicht", "Totaalproducten en lage voorraad."),
        ("Producten", "Productbeheer", "Zoeken, filteren, toevoegen en bewerken."),
        ("Orders", "Orderhistorie", "Klantorders en datums terugvinden."),
        ("Klanten", "Klantenbeheer", "Klanten bekijken en beheren."),
    ]
    for i, (k, h, p) in enumerate(cards):
        cx = x + 48 + i * 304
        parts += [
            rect(cx, y + 510, 280, 190, COLORS["surface"], COLORS["border"], 8),
            text(cx + 22, y + 552, k.upper(), 12, 800, COLORS["heading"]),
            text(cx + 22, y + 592, h, 21, 800, COLORS["ink"]),
            text(cx + 22, y + 632, p, 14, 400, COLORS["muted"]),
        ]
    parts += [
        rect(x + 48, y + 750, 1188, 90, COLORS["surface_alt"], COLORS["border"], 8),
        text(x + 76, y + 792, "Gebruikelijke workflow", 20, 800, COLORS["ink"]),
        text(x + 76, y + 820, "Dashboard controleren -> Producten bijwerken -> Orders raadplegen", 15, 400, COLORS["muted"]),
        *button(x + 1040, y + 774, "Nieuw product", "success", 160),
    ]
    return parts


def dashboard_frame(x, y):
    parts = [*frame_label(x, y - 22, "02 Admin dashboard", "Status en snelle acties"), *shell(x, y, "Admin dashboard", "Beheer producten en houd voorraad snel in de gaten.", "Admin")]
    parts += [
        *button(x + 1014, y + 104, "Producten bekijken", "primary", 180),
        *button(x + 1208, y + 104, "Nieuw product", "outline", 130),
    ]
    stats = [("Totaal producten", "3", COLORS["surface"]), ("Lage voorraad", "2", "#FFF8E1")]
    for i, (label, value, fill) in enumerate(stats):
        cx = x + 48 + i * 318
        parts += [
            rect(cx, y + 210, 290, 150, fill, COLORS["border"], 8),
            text(cx + 26, y + 258, label, 15, 700, COLORS["muted"]),
            text(cx + 26, y + 320, value, 52, 800, COLORS["ink"]),
        ]
    parts += [
        rect(x + 48, y + 410, 1188, 330, COLORS["surface"], COLORS["border"], 8),
        text(x + 76, y + 462, "Beheertaken", 24, 800, COLORS["ink"]),
        text(x + 76, y + 504, "1. Controleer lage voorraad", 18, 700, COLORS["heading"]),
        text(x + 76, y + 542, "2. Werk productinformatie bij", 18, 700, COLORS["heading"]),
        text(x + 76, y + 580, "3. Raadpleeg orders bij klantvragen", 18, 700, COLORS["heading"]),
        rect(x + 720, y + 456, 430, 210, COLORS["surface_alt"], COLORS["border"], 8),
        text(x + 748, y + 504, "HCI onderbouwing", 22, 800, COLORS["ink"]),
        text(x + 748, y + 548, "Het dashboard toont alleen de belangrijkste signalen.", 15, 400, COLORS["muted"]),
        text(x + 748, y + 582, "Actieknoppen sluiten aan op de meest voorkomende taken.", 15, 400, COLORS["muted"]),
    ]
    return parts


def product_frame(x, y):
    parts = [*frame_label(x, y - 22, "03 Productbeheer", "Zoeken, filteren en CRUD"), *shell(x, y, "Productbeheer", "Zoek, filter en beheer producten van Matrix Inc.", "Products")]
    parts += [
        *button(x + 1086, y + 104, "Nieuw product", "primary", 150),
        rect(x + 48, y + 194, 1188, 86, COLORS["surface_alt"], COLORS["border"], 8),
        text(x + 76, y + 232, "Productnaam", 13, 700, COLORS["ink"]),
        rect(x + 76, y + 246, 330, 38, COLORS["surface"], COLORS["border"], 6),
        text(x + 94, y + 271, "Zoek op naam", 13, 400, COLORS["muted"]),
        rect(x + 450, y + 246, 18, 18, COLORS["surface"], COLORS["border"], 4),
        text(x + 478, y + 261, "Alleen lage voorraad", 14, 500, COLORS["ink"]),
        *button(x + 690, y + 240, "Filter", "primary", 100),
        *button(x + 804, y + 240, "Reset", "outline", 90),
        text(x + 48, y + 326, "3 resultaten  |  2 van 3 producten hebben lage voorraad", 14, 600, COLORS["muted"]),
        rect(x + 48, y + 360, 1188, 360, COLORS["surface"], COLORS["border"], 8),
    ]
    headers = ["Naam", "Omschrijving", "Prijs", "Voorraad", "Onderdelen", "Acties"]
    widths = [190, 420, 130, 130, 120, 198]
    cx = x + 48
    parts.append(rect(cx, y + 360, 1188, 50, COLORS["hero"], None, 8))
    for h, w in zip(headers, widths):
        parts.append(text(cx + 20, y + 392, h, 13, 800, COLORS["heading"]))
        cx += w
    rows = [
        ("Nebuchadnezzar", "Schip uit de Matrix wereld", "EUR 10000", "3 laag", "0"),
        ("Jack-in Chair", "Stoel om in te pluggen", "EUR 500,50", "12", "0"),
        ("EMP Device", "Wapentuig op schepen", "EUR 129,99", "4 laag", "0"),
    ]
    for r, row in enumerate(rows):
        yy = y + 410 + r * 86
        parts.append(line(x + 48, yy, x + 1236, yy))
        cx = x + 48
        for idx, (value, w) in enumerate(zip(row, widths[:-1])):
            if idx == 3 and "laag" in value:
                parts.append(rect(cx + 18, yy + 24, 72, 28, COLORS["warning"], None, 14))
                parts.append(text(cx + 54, yy + 43, value, 12, 800, COLORS["ink"], "middle"))
            else:
                parts.append(text(cx + 20, yy + 48, value, 13, 500, COLORS["ink"] if idx == 0 else COLORS["muted"]))
            cx += w
        parts += [
            *button(x + 1036, yy + 22, "Details", "outline", 76),
            *button(x + 1120, yy + 22, "Edit", "outline", 56),
            *button(x + 1184, yy + 22, "Delete", "danger", 70),
        ]
    return parts


def orders_frame(x, y):
    parts = [*frame_label(x, y - 22, "04 Orderhistorie", "Filteren op klant, datum en zoekterm"), *shell(x, y, "Orderhistorie", "Alle orders van alle klanten in een overzicht.", "Orders")]
    parts += [
        rect(x + 1070, y + 106, 166, 34, "#6C757D", None, 17),
        text(x + 1153, y + 128, "4 orders", 14, 800, "#FFFFFF", "middle"),
        rect(x + 48, y + 194, 1188, 100, COLORS["surface_alt"], COLORS["border"], 8),
        text(x + 76, y + 232, "Klant", 13, 700, COLORS["ink"]),
        rect(x + 76, y + 246, 210, 38, COLORS["surface"], COLORS["border"], 6),
        text(x + 94, y + 271, "Alle klanten", 13, 500, COLORS["muted"]),
        text(x + 312, y + 232, "Zoeken", 13, 700, COLORS["ink"]),
        rect(x + 312, y + 246, 260, 38, COLORS["surface"], COLORS["border"], 6),
        text(x + 330, y + 271, "Ordernr, klant of adres", 13, 500, COLORS["muted"]),
        text(x + 598, y + 232, "Vanaf", 13, 700, COLORS["ink"]),
        rect(x + 598, y + 246, 150, 38, COLORS["surface"], COLORS["border"], 6),
        text(x + 772, y + 232, "Tot en met", 13, 700, COLORS["ink"]),
        rect(x + 772, y + 246, 150, 38, COLORS["surface"], COLORS["border"], 6),
        *button(x + 962, y + 242, "Filter", "primary", 92),
        *button(x + 1066, y + 242, "Reset", "outline", 90),
        rect(x + 48, y + 342, 1188, 360, COLORS["surface"], COLORS["border"], 8),
    ]
    headers = ["Ordernummer", "Orderdatum", "Klant", "Adres", "Status", "Producten"]
    widths = [170, 170, 180, 260, 150, 258]
    cx = x + 48
    parts.append(rect(cx, y + 342, 1188, 50, COLORS["hero"], None, 8))
    for h, w in zip(headers, widths):
        parts.append(text(cx + 20, y + 374, h, 13, 800, COLORS["heading"]))
        cx += w
    rows = [
        ("#4", "01-03-2021", "Trinity", "789 Pine St", "Actief", "Geen producten gekoppeld"),
        ("#3", "01-02-2021", "Morpheus", "456 Oak St", "Actief", "Geen producten gekoppeld"),
        ("#2", "01-02-2021", "Neo", "123 Elm St", "Actief", "Geen producten gekoppeld"),
        ("#1", "01-01-2021", "Neo", "123 Elm St", "Actief", "Geen producten gekoppeld"),
    ]
    for r, row in enumerate(rows):
        yy = y + 392 + r * 70
        parts.append(line(x + 48, yy, x + 1236, yy))
        cx = x + 48
        for idx, (value, w) in enumerate(zip(row, widths)):
            if idx == 4:
                parts.append(rect(cx + 18, yy + 20, 70, 28, COLORS["success"], None, 14))
                parts.append(text(cx + 53, yy + 39, "Actief", 12, 800, "#FFFFFF", "middle"))
            else:
                parts.append(text(cx + 20, yy + 42, value, 13, 500, COLORS["ink"] if idx in (0, 2) else COLORS["muted"]))
            cx += w
    return parts


def component_frame(x, y):
    parts = [*frame_label(x, y - 22, "05 Componenten en tokens", "Basis voor Figma componenten")]
    parts += [
        rect(x, y, 1366, 620, COLORS["page"], COLORS["border"], 8),
        text(x + 48, y + 76, "Design system basis", 34, 800, COLORS["ink"]),
        text(x + 48, y + 112, "Kleuren, typografie, knoppen, kaarten, filters en tabelpatronen.", 16, 400, COLORS["muted"]),
        text(x + 48, y + 170, "Kleuren", 22, 800, COLORS["ink"]),
    ]
    swatches = [("Ink", "ink"), ("Primary", "primary"), ("Success", "success"), ("Warning", "warning"), ("Danger", "danger"), ("Surface", "surface_alt")]
    for i, (label, key) in enumerate(swatches):
        cx = x + 48 + i * 132
        parts += [
            rect(cx, y + 196, 96, 64, COLORS[key], COLORS["border"], 8),
            text(cx, y + 286, label, 13, 700, COLORS["ink"]),
            text(cx, y + 306, COLORS[key], 12, 500, COLORS["muted"]),
        ]
    parts += [
        text(x + 48, y + 370, "Knoppen", 22, 800, COLORS["ink"]),
        *button(x + 48, y + 396, "Primair", "primary", 118),
        *button(x + 182, y + 396, "Secundair", "outline", 128),
        *button(x + 326, y + 396, "Opslaan", "success", 112),
        *button(x + 454, y + 396, "Verwijderen", "danger", 138),
        text(x + 720, y + 370, "Formulier veld", 22, 800, COLORS["ink"]),
        text(x + 720, y + 414, "Productnaam", 13, 700, COLORS["ink"]),
        rect(x + 720, y + 428, 360, 42, COLORS["surface"], COLORS["border"], 6),
        text(x + 740, y + 455, "Nebuchadnezzar", 14, 500, COLORS["ink"]),
        text(x + 720, y + 502, "Foutmelding: Productnaam is verplicht.", 13, 600, COLORS["danger"]),
    ]
    return parts


def build_svg():
    width, height = 2860, 2140
    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        rect(0, 0, width, height, "#EEF2F6", None, 0),
        text(48, 58, "Matrix Inc. Back-office - Figma-ready interface design", 28, 800, COLORS["ink"]),
        text(48, 88, "Importeer deze SVG in Figma. Elk genummerd vlak is bedoeld als apart frame.", 15, 500, COLORS["muted"]),
    ]
    parts += home_frame(48, 150)
    parts += dashboard_frame(1494, 150)
    parts += product_frame(48, 1110)
    parts += orders_frame(1494, 1110)
    parts += component_frame(48, 2040)
    parts.append("</svg>")
    CANVAS.write_text("\n".join(parts), encoding="utf-8")

    for name, fn in [
        ("01-home.svg", home_frame),
        ("02-admin-dashboard.svg", dashboard_frame),
        ("03-productbeheer.svg", product_frame),
        ("04-orderhistorie.svg", orders_frame),
        ("05-componenten.svg", component_frame),
    ]:
        frame = [
            '<svg xmlns="http://www.w3.org/2000/svg" width="1462" height="982" viewBox="0 0 1462 982">',
            rect(0, 0, 1462, 982, "#EEF2F6", None, 0),
        ]
        frame += fn(48, 72)
        frame.append("</svg>")
        (OUT / name).write_text("\n".join(frame), encoding="utf-8")


def build_markdown():
    TOKENS.write_text(json.dumps(TOKENS_DATA, indent=2, ensure_ascii=False), encoding="utf-8")

    GUIDE.write_text(
        """# Figma import guide - Matrix Inc. Back-office

## Bestanden

- `matrix-inc-backoffice-figma.svg`: volledig canvas met alle schermen en componenten.
- `01-home.svg`: homepagina.
- `02-admin-dashboard.svg`: admin dashboard.
- `03-productbeheer.svg`: productbeheer.
- `04-orderhistorie.svg`: orderhistorie.
- `05-componenten.svg`: componenten en design tokens.
- `design-tokens.json`: kleuren, typografie, spacing en HCI-principes.
- `interface-design-onderbouwing.md`: Nederlandse onderbouwing.
- `Matrix_Inc_Figma_Interface_Design.docx`: Word-versie van de onderbouwing.

## Importeren in Figma

1. Open Figma en maak een nieuw designbestand.
2. Sleep `matrix-inc-backoffice-figma.svg` op het canvas.
3. Selecteer de geimporteerde frames en groepeer per scherm.
4. Maak componenten van knoppen, kaarten, badges, filtervelden en tabelrijen.
5. Gebruik `design-tokens.json` als basis voor lokale styles in Figma.

## Prototype flow

Home -> Admin dashboard -> Productbeheer -> Product toevoegen

Home -> Orderhistorie

Home -> Klantenbeheer
""",
        encoding="utf-8",
    )

    RATIONALE_MD.write_text(
        """# Interface design onderbouwing - Matrix Inc. Back-office

## Doel

Het ontwerp ondersteunt een admin back-office voor Matrix Inc. De gebruiker moet snel kunnen navigeren naar productbeheer, orderhistorie, klantenbeheer en dashboardinformatie.

## Ontwerpkeuzes

### Homepagina

De homepagina is ontworpen als startpunt voor dagelijks beheer. De primaire acties staan bovenaan: dashboard openen en productbeheer. Daaronder staan vier navigatiekaarten voor de belangrijkste routes.

### Admin dashboard

Het dashboard toont alleen de belangrijkste signalen: totaal aantal producten en lage voorraad. Dit voorkomt onnodige cognitieve belasting en helpt de gebruiker direct prioriteiten te zien.

### Productbeheer

Productbeheer gebruikt een tabel omdat producten vergelijkbare eigenschappen hebben: naam, omschrijving, prijs, voorraad en acties. De zoekbalk en lage-voorraadfilter staan boven de tabel, zodat de gebruiker eerst kan beperken en daarna kan handelen.

### Orderhistorie

Orderhistorie gebruikt filters voor klant, zoekterm en datumbereik. Dit sluit aan op de taak: snel een order terugvinden bij een klantvraag.

## HCI-principes

- Zichtbaarheid: belangrijke routes staan zichtbaar op de homepagina.
- Consistentie: knoppen, kaarten en tabellen gebruiken dezelfde visuele stijl.
- Feedback: status en lage voorraad worden met badges weergegeven.
- Foutpreventie: verwijderen blijft een aparte bevestigingsstap in de applicatie.
- Herkenbaarheid: tabellen en formulieren volgen bekende admin patronen.

## Design system

De interface gebruikt een rustige zakelijke kleurstelling met blauw als primaire actiekleur. De achtergrond is licht, kaarten zijn wit en randen zijn subtiel. Daardoor blijft de applicatie overzichtelijk en geschikt voor herhaald administratief gebruik.

## Responsive gedrag

Op desktop staan kaarten en filters naast elkaar. Op mobiel stapelen de elementen onder elkaar, zodat knoppen en invoervelden goed bereikbaar blijven.
""",
        encoding="utf-8",
    )


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2

    title = doc.add_paragraph()
    run = title.add_run("Interface design onderbouwing - Matrix Inc.")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph("Figma-ready ontwerp voor de ASP.NET Core MVC admin back-office.")
    subtitle.runs[0].font.color.rgb = RGBColor(95, 107, 122)

    sections = [
        ("Doel", "Het ontwerp ondersteunt een admin back-office waarin gebruikers snel naar productbeheer, orderhistorie, klantenbeheer en dashboardinformatie kunnen navigeren."),
        ("Homepagina", "De homepagina is het centrale startpunt. De hero-sectie geeft context en de navigatiekaarten maken de belangrijkste taken direct bereikbaar."),
        ("Admin dashboard", "Het dashboard toont alleen kerninformatie: totaal aantal producten en lage voorraad. Dit helpt de gebruiker snel prioriteiten te bepalen."),
        ("Productbeheer", "Het productoverzicht gebruikt een scanbare tabel met zoekfunctie, lage-voorraadfilter en duidelijke actieknoppen voor details, bewerken en verwijderen."),
        ("Orderhistorie", "De orderpagina biedt filters op klant, zoekterm en datumbereik, passend bij het terugvinden van orders bij klantvragen."),
        ("HCI-principes", "Het ontwerp gebruikt zichtbaarheid, consistentie, feedback, foutpreventie en herkenbare admin patronen."),
        ("Responsive gedrag", "Op desktop worden elementen naast elkaar geplaatst. Op mobiel stapelen kaarten, filters en knoppen onder elkaar."),
    ]
    for heading, body in sections:
        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = RGBColor(46, 116, 181)
        doc.add_paragraph(body)

    doc.add_heading("Bestanden in dit pakket", level=1)
    for item in [
        "matrix-inc-backoffice-figma.svg",
        "01-home.svg",
        "02-admin-dashboard.svg",
        "03-productbeheer.svg",
        "04-orderhistorie.svg",
        "05-componenten.svg",
        "design-tokens.json",
        "figma-import-guide.md",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(RATIONALE_DOCX)


if __name__ == "__main__":
    build_svg()
    build_markdown()
    build_docx()
    print(OUT)
